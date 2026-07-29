from __future__ import annotations

import shutil
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.opc.constants import RELATIONSHIP_TYPE

import build_submission_report as base


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "report" / "submission"
OUT = ROOT / "report" / "final_submission"
STEM = "Peleg_Shpitzer_Razi_Mreeh_Final_Report"
SOURCE_MD = SOURCE_DIR / "Peleg_Shpitzer_Razi_Mreeh_VGGT_Seminar_Report.md"
TARGET_MD = OUT / f"{STEM}.md"
TARGET_DOCX = OUT / f"{STEM}.docx"
FIGURES = OUT / "figures"
HEADER = "SEMINAR IN DEEP LEARNING FOR COMPUTER VISION | VGGT EVALUATION"

TITLE_BLOCK = """# Final Project
## Seminar in Deep Learning for Solving Computer Vision Problems

# Reproducing and Evaluating VGGT:
## View-Count Scaling, Confidence Calibration, and Bounded Domain Adaptation

**A Controlled Study on ETH3D and TartanAir**

**Peleg Shpitzer**  
**Razi Mreeh**
"""


def prepare_markdown() -> list[str]:
    source = SOURCE_MD.read_text(encoding="utf-8")
    marker = "<!-- PAGE BREAK -->"
    _, remainder = source.split(marker, 1)
    final_text = f"{TITLE_BLOCK}\n{marker}{remainder}"
    adaptation_figure = "![Adaptation validation](figures/adaptation_validation.png)"
    final_text = final_text.replace(
        adaptation_figure,
        "*Table 2. Bounded adaptation configuration and resource record.*\n\n"
        + adaptation_figure,
        1,
    )
    OUT.mkdir(parents=True, exist_ok=True)
    TARGET_MD.write_text(final_text, encoding="utf-8")
    return final_text.splitlines()


def prepare_figures() -> None:
    if FIGURES.exists():
        shutil.rmtree(FIGURES)
    shutil.copytree(SOURCE_DIR / "figures", FIGURES)


def configure_final(doc: Document) -> None:
    base.configure(doc)
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.first_page_header.paragraphs[0].text = ""
    header = section.header.paragraphs[0]
    header.text = HEADER
    for run in header.runs:
        base.set_font(run, 8.5, True, base.MUTED)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(64)
    p.paragraph_format.space_after = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_font(p.add_run("Final Project"), 11, True, base.GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    base.set_font(
        p.add_run("Seminar in Deep Learning for Solving Computer Vision Problems"),
        13,
        True,
        base.BLUE,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    title = (
        "Reproducing and Evaluating VGGT:\n"
        "View-Count Scaling, Confidence Calibration,\n"
        "and Bounded Domain Adaptation"
    )
    base.set_font(p.add_run(title), 25, True, base.NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(58)
    base.set_font(
        p.add_run("A Controlled Study on ETH3D and TartanAir"),
        14,
        color=base.BLUE,
    )

    for author in ("Peleg Shpitzer", "Razi Mreeh"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        base.set_font(p.add_run(author), 14, True, base.NAVY)


def make_urls_clickable(doc: Document) -> None:
    pattern = re.compile(r"(https?://[^\s)]+)")
    for paragraph in doc.paragraphs:
        for run in list(paragraph.runs):
            match = pattern.search(run.text)
            if not match:
                continue
            before, url, after = run.text[: match.start()], match.group(1), run.text[match.end() :]
            trailing = ""
            while url.endswith((".", ",", ";")):
                trailing = url[-1] + trailing
                url = url[:-1]
            run.text = before
            relationship = doc.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), relationship)
            linked_run = OxmlElement("w:r")
            props = OxmlElement("w:rPr")
            color = OxmlElement("w:color")
            color.set(qn("w:val"), "2E74B5")
            underline = OxmlElement("w:u")
            underline.set(qn("w:val"), "single")
            props.extend((color, underline))
            text = OxmlElement("w:t")
            text.text = url
            linked_run.extend((props, text))
            hyperlink.append(linked_run)
            run._r.addnext(hyperlink)
            if after or trailing:
                suffix = OxmlElement("w:r")
                suffix_text = OxmlElement("w:t")
                suffix_text.text = trailing + after
                suffix.append(suffix_text)
                hyperlink.addnext(suffix)


def main() -> None:
    lines = prepare_markdown()
    prepare_figures()
    base.FIGURES = FIGURES

    doc = Document()
    configure_final(doc)
    add_cover(doc)
    marker_index = lines.index("<!-- PAGE BREAK -->")
    base.body(doc, lines[marker_index:])
    make_urls_clickable(doc)
    doc.save(TARGET_DOCX)
    print(TARGET_DOCX)


if __name__ == "__main__":
    main()
