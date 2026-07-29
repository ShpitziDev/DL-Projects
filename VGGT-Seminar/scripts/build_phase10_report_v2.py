"""Build the visual-first V2 report from saved Phase 3-9 artifacts only.

This script never imports VGGT or Torch, loads no checkpoint, and performs no inference.
"""
from pathlib import Path
import csv, json, re, shutil
from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT=Path(__file__).resolve().parents[1]; V2=ROOT/'report/v2'; FIG=V2/'figures'; DIA=V2/'diagrams'; TAB=V2/'tables'; SUP=V2/'supplementary'; BUILD=V2/'build'
DATA=ROOT/'report/data/view_count_results.csv'; COUNTS=[2,4,6,8,10]
NAVY='#14283D'; BLUE='#2B6CB0'; CYAN='#20A4B8'; ORANGE='#E97932'; PALE='#EEF4F8'; INK='#1C2733'; MUTED='#5D6B78'

def font(size,bold=False):
    try:return ImageFont.truetype('arialbd.ttf' if bold else 'arial.ttf',size)
    except OSError:return ImageFont.load_default()
def label(draw,xy,text,size=24,color=INK,bold=False,anchor=None): draw.text(xy,text,font=font(size,bold),fill=color,anchor=anchor)
def fit(path,size):
    im=Image.open(path).convert('RGB'); return ImageOps.fit(im,size,Image.Resampling.LANCZOS)
def save(im,path): path.parent.mkdir(parents=True,exist_ok=True); im.save(path,quality=94)
def pred(scene,count):
    if scene=='delivery_area' and count==2: base=ROOT/'outputs/predictions/phase6_2_eth3d_overlap_smoke/delivery_area/S2_overlap_aware_nested_original'
    elif scene=='delivery_area': base=ROOT/f'outputs/predictions/phase7_eth3d_view_count/delivery_area/S{count}_overlap_aware_nested_original'
    else: base=ROOT/f'outputs/predictions/phase8_eth3d_view_count/courtyard/S{count}_overlap_aware_nested_original'
    return base/'visualizations'
def scene_image(scene,name): return ROOT/f'local_assets/datasets/eth3d/{scene}/images/dslr_images_undistorted/{name}'

def hero():
    im=Image.new('RGB',(1800,920),NAVY); d=ImageDraw.Draw(im)
    tiles=[scene_image('delivery_area','DSC_0675.JPG'),scene_image('courtyard','DSC_0286.JPG'),pred('courtyard',10)/'depth_view0.png',FIG/'delivery_area_s10_pointcloud_hero.png']
    x=[40,485,930,1375]
    for i,p in enumerate(tiles):
        t=fit(p,(385,580)); im.paste(t,(x[i],120)); d.rounded_rectangle((x[i]-3,117,x[i]+388,703),radius=14,outline='#8ED6E1',width=4)
    for a,b in zip(x,x[1:]): d.line((a+390,410,b-8,410),fill='#DDEAF1',width=5); d.polygon([(b-8,410),(b-30,397),(b-30,423)],fill='#DDEAF1')
    for xx,t in zip(x,['INPUT / DELIVERY','INPUT / COURTYARD','PREDICTED DEPTH','RECONSTRUCTED GEOMETRY']): label(d,(xx+192,755),t,22,'#EAF5FA',True,'mm')
    label(d,(900,58),'FROM OVERLAPPING IMAGES TO VISUAL GEOMETRY',32,'#FFFFFF',True,'mm'); save(im,DIA/'hero_visual.png')

def glance(rows):
    im=Image.new('RGB',(1800,1100),'white'); d=ImageDraw.Draw(im); label(d,(90,70),'STUDY AT A GLANCE',50,NAVY,True)
    cards=[('INPUT','2 / 4 / 6 / 8 / 10','overlapping views'),('MODEL','VGGT-1B','feed-forward'),('SCENES','2','delivery_area + courtyard'),('CONFIGURATIONS','10','all outputs finite')]
    for i,(k,v,s) in enumerate(cards):
        x=90+i*420; d.rounded_rectangle((x,155,x+370,390),radius=24,fill=PALE,outline='#C7D8E5',width=3); label(d,(x+24,185),k,18,BLUE,True); label(d,(x+185,270),v,42,NAVY,True,'mm'); label(d,(x+185,340),s,20,MUTED,False,'mm')
    findings=[('~1 s','synchronized inference'),('5.26 → 6.60 GiB','allocated VRAM'),('NON-MONOTONIC','model confidence'),('8–10 VIEWS','diminishing visible gains'),('NO CLAIM','of quantitative accuracy')]
    for i,(v,s) in enumerate(findings):
        x=90+(i%3)*560; y=485+(i//3)*240; d.rounded_rectangle((x,y,x+510,y+190),radius=18,fill=NAVY if i<4 else '#6B3E22'); label(d,(x+255,y+70),v,32,'white',True,'mm'); label(d,(x+255,y+132),s,19,'#DCE8EF',False,'mm')
    label(d,(90,1015),'Interpretation boundary: confidence is model output; camera and geometry units are arbitrary and unaligned.',22,'#7A4B00',True); save(im,DIA/'study_at_a_glance.png')

def flow_diagram(path,title,steps,branches=False):
    im=Image.new('RGB',(1800,980),'white'); d=ImageDraw.Draw(im); label(d,(80,55),title,44,NAVY,True)
    if not branches:
        y=145
        for i,(head,sub) in enumerate(steps):
            d.rounded_rectangle((300,y,1500,y+95),radius=22,fill=PALE if i%2==0 else '#E7F6F7',outline=BLUE,width=3); label(d,(900,y+36),head,26,NAVY,True,'mm'); label(d,(900,y+70),sub,17,MUTED,False,'mm')
            if i<len(steps)-1: d.line((900,y+95,900,y+125),fill=ORANGE,width=5); d.polygon([(900,y+132),(888,y+112),(912,y+112)],fill=ORANGE)
            y+=125
    else:
        label(d,(900,145),'MULTI-VIEW RGB IMAGES',28,NAVY,True,'mm'); d.rounded_rectangle((620,110,1180,190),radius=18,fill=PALE,outline=BLUE,width=3)
        for y,text,sub in [(270,'DINOv2 IMAGE-TOKEN ENCODER','per-image visual tokens'),(420,'ALTERNATING ATTENTION AGGREGATOR','frame-local ↔ cross-view global exchange')]:
            d.line((900,y-80,900,y-25),fill=ORANGE,width=5); d.polygon([(900,y-18),(888,y-38),(912,y-38)],fill=ORANGE); d.rounded_rectangle((450,y-20,1350,y+80),radius=20,fill='#E7F6F7',outline=CYAN,width=4); label(d,(900,y+15),text,27,NAVY,True,'mm'); label(d,(900,y+52),sub,18,MUTED,False,'mm')
        label(d,(900,550),'FIRST-CAMERA REFERENCE FRAME',18,'#7A4B00',True,'mm'); d.line((900,500,900,585),fill=ORANGE,width=5)
        heads=[('CAMERA','pose + intrinsics'),('DEPTH','map + confidence'),('POINT MAP','3D + confidence'),('TRACKING','tracks + visibility')]
        for i,(h,s) in enumerate(heads):
            x=70+i*430; d.rounded_rectangle((x,620,x+390,780),radius=22,fill=NAVY,outline=NAVY,width=3); label(d,(x+195,675),h,25,'white',True,'mm'); label(d,(x+195,725),s,18,'#CFE5EF',False,'mm')
        label(d,(900,875),'Feed-forward inference · unified output heads · no iterative reconstruction loop',23,BLUE,True,'mm')
    save(im,path)

def endpoint_figure():
    im=Image.new('RGB',(1800,940),'white'); d=ImageDraw.Draw(im); label(d,(70,55),'WHY OVERLAP-AWARE SELECTION MATTERS',42,NAVY,True)
    groups=[('ENDPOINTS [0, 43]','Limited shared visibility',['DSC_0675.JPG','DSC_0718.JPG'],'#B84B3A'),('OVERLAP-AWARE [0, 6]','Doorway · pillar · ceiling · wall retained',['DSC_0675.JPG','DSC_0681.JPG'],CYAN)]
    for gi,(head,sub,names,color) in enumerate(groups):
        y=150+gi*380; label(d,(70,y),head,26,color,True)
        for j,name in enumerate(names):
            x=70+j*835; tile=fit(scene_image('delivery_area',name),(770,280)); im.paste(tile,(x,y+55)); d.rectangle((x,y+55,x+770,y+335),outline=color,width=5); label(d,(x+385,y+355),name,18,MUTED,False,'mm')
        label(d,(1700,y+195),sub,18,color,True,'rm')
    save(im,FIG/'overlap_pair_comparison.png')

def nested(scene, rows):
    im=Image.new('RGB',(1800,1250),'white'); d=ImageDraw.Draw(im); label(d,(70,55),f'{scene.replace("_"," ").upper()} · NESTED VIEW PROTOCOL',42,NAVY,True); label(d,(70,110),'S2 ⊂ S4 ⊂ S6 ⊂ S8 ⊂ S10  |  colored borders mark newly added views',22,MUTED)
    previous=set()
    for ri,count in enumerate(COUNTS):
        row=next(r for r in rows if r['scene']==scene and int(r['view_count'])==count); names=json.loads(row['filenames']); y=175+ri*205; label(d,(70,y+65),f'S{count}',28,BLUE,True)
        for j,name in enumerate(names):
            tile=fit(scene_image(scene,name),(145,150)); x=175+j*158; im.paste(tile,(x,y)); color=ORANGE if name not in previous else '#B9C7D3'; d.rectangle((x,y,x+145,y+150),outline=color,width=5)
        label(d,(175,y+175),', '.join(str(v) for v in json.loads(row['indices'])),17,MUTED)
        previous.update(names)
    save(im,FIG/f'nested_{scene}.png')

def chart(rows,kind,path,title,note):
    im=Image.new('RGB',(1600,920),'white'); d=ImageDraw.Draw(im); label(d,(80,55),title,38,NAVY,True); label(d,(80,105),note,19,'#7A4B00',True)
    left,top,right,bottom=150,180,1510,760; d.line((left,top,left,bottom),fill='#506070',width=3); d.line((left,bottom,right,bottom),fill='#506070',width=3)
    series=[]
    if kind=='time': series=[('delivery_area','total_seconds',BLUE),('courtyard','total_seconds',ORANGE)]
    elif kind=='memory': series=[('allocated','peak_allocated_gib',BLUE),('reserved','peak_reserved_gib',ORANGE)]
    elif kind=='confidence': series=[('delivery depth','depth_confidence_mean',BLUE),('courtyard depth','depth_confidence_mean',ORANGE),('delivery point','point_confidence_mean','#1A8F75'),('courtyard point','point_confidence_mean','#8B5FBF')]
    else: series=[('delivery_area','camera_path_length',BLUE),('courtyard','camera_path_length',ORANGE)]
    vals=[]
    for name,key,color in series:
        if kind=='memory': vals.append((name,[float(next(r for r in rows if r['scene']=='delivery_area' and int(r['view_count'])==c)[key]) for c in COUNTS],color))
        else:
            scene='delivery_area' if name.startswith('delivery') else 'courtyard'; vals.append((name,[float(next(r for r in rows if r['scene']==scene and int(r['view_count'])==c)[key]) for c in COUNTS],color))
    ymax=max(max(v) for _,v,_ in vals)*1.12; ymin=0
    for i in range(6): y=bottom-(bottom-top)*i/5; d.line((left,y,right,y),fill='#E3E8ED',width=2); label(d,(left-18,y),f'{ymax*i/5:.1f}',16,MUTED,False,'rm')
    for ci,c in enumerate(COUNTS): x=left+(right-left)*ci/4; label(d,(x,bottom+36),str(c),18,INK,False,'mm')
    for si,(name,values,color) in enumerate(vals):
        pts=[]
        for ci,v in enumerate(values): x=left+(right-left)*ci/4; y=bottom-(v-ymin)/(ymax-ymin)*(bottom-top); pts.append((x,y))
        d.line(pts,fill=color,width=7,joint='curve')
        for x,y in pts: d.ellipse((x-9,y-9,x+9,y+9),fill='white',outline=color,width=5)
        lx=180+(si%2)*560; ly=835+(si//2)*35; d.line((lx,ly,lx+50,ly),fill=color,width=7); label(d,(lx+65,ly),name.replace('_',' '),18,INK,False,'lm')
    label(d,(830,805),'Input views',20,INK,True,'mm'); save(im,path)

def gallery(scene,metric,out,title):
    im=Image.new('RGB',(1800,720),NAVY); d=ImageDraw.Draw(im); label(d,(55,45),title,36,'white',True)
    for i,c in enumerate(COUNTS):
        p=pred(scene,c)/f'{metric}_view0.png'; tile=fit(p,(325,500)); x=45+i*350; im.paste(tile,(x,120)); d.rectangle((x,120,x+325,620),outline='#7FD1DD',width=3); label(d,(x+162,660),f'S{c}',24,'white',True,'mm')
    save(im,FIG/out)

def output_diversity():
    scene='courtyard'; c=10; sources=[scene_image(scene,'DSC_0286.JPG'),pred(scene,c)/'depth_view0.png',pred(scene,c)/'depth_confidence_view0.png',pred(scene,c)/'point_confidence_view0.png',pred(scene,c)/'camera_centers.png',FIG/'courtyard_s10_pointcloud_hero.png']; names=['Input view','Depth','Depth confidence','Point confidence','Cameras','Geometry']
    im=Image.new('RGB',(1800,760),'white'); d=ImageDraw.Draw(im); label(d,(50,40),'ONE MODEL · MULTIPLE GEOMETRIC OUTPUTS',38,NAVY,True)
    for i,(p,n) in enumerate(zip(sources,names)):
        x=35+i*292; tile=fit(p,(270,540)); im.paste(tile,(x,115)); d.rectangle((x,115,x+270,655),outline=BLUE,width=3); label(d,(x+135,695),n,19,INK,True,'mm')
    save(im,FIG/'output_diversity.png')

def md_report():
    text="""# Reproducing VGGT: A Controlled Study of View-Count Scaling with Overlap-Aware ETH3D Inputs

Peleg Shpitzer · Course / Instructor · University of Haifa · July 2026

## Abstract
Visual Geometry Grounded Transformer (VGGT) predicts cameras, depth, point maps, confidence, and tracks from one or more images in a single feed-forward model. We reproduced the maintained official implementation and public VGGT-1B checkpoint locally, then designed a controlled study of how operational behavior changes with 2, 4, 6, 8, and 10 overlapping views. Two calibrated ETH3D scenes, delivery_area and courtyard, were sampled with a deterministic overlap-aware protocol combining pose and feature evidence. Across ten configurations, all required outputs remained finite. Allocated memory rose from about 5.26 to 6.60 GiB, synchronized inference stayed near one second, and complete processing time grew with the amount of per-view postprocessing. Model confidence generally increased but was not monotonic, while visible gains appeared to diminish around eight to ten views. Dominant scene structure remained qualitatively coherent. Because predictions were not aligned to ETH3D coordinates, this report makes no quantitative reconstruction-accuracy claim. The results support reliable local operation and a defensible view-count methodology, while identifying aligned evaluation as the main next scientific step.

## Study at a Glance
Ten saved configurations span two scenes and five nested view counts. The design isolates view count while retaining earlier frames in every larger set.

## Introduction
VGGT asks whether a shared transformer can replace much of a conventional multi-stage geometry pipeline with direct prediction [1]. This project tests that proposition at the level appropriate for a seminar reproduction: operational correctness, a controlled input study, resource measurements, and bounded qualitative analysis. The research question is how runtime, memory, confidence, predicted camera structure, and visible completeness change as overlapping input views increase from two to ten.

## VGGT in Brief
Multi-view RGB images are tokenized by a DINOv2 encoder [5]. An aggregator alternates frame-local attention with global cross-view attention, after which dedicated heads predict cameras, depth, point maps, confidence, and tracks. Predictions use the first camera as a reference. The conceptual diagram in this report is original and based on the architecture description in VGGT [1], not a reproduction of its figure.

## Reproduction Setup
The maintained official repository was pinned at `a288dd0f14786c93483e45524328726ab7b1b4ce`; the public checkpoint hash was `d15bf50a8615c8225ed48b51ea5cac673d82442ec0309036df555a053253afe0`. Inference used Python 3.11.15, PyTorch 2.13.0+cu130, CUDA 13.0, an RTX 5080, BF16 autocast, and disabled Flash SDPA. The official example and all documented output heads were validated. This establishes operational reproduction, not independent reproduction of every benchmark result reported by the authors.

## Experimental Methodology
ETH3D provides calibrated high-resolution imagery and laser scans [2]. We selected delivery_area and courtyard as complementary indoor and outdoor cases. Laser geometry was deliberately not used for scoring because coordinate conventions and similarity alignment must be validated first. The independent variable was view count; model, checkpoint, preprocessing, precision, hardware, ordering, postprocessing, and the strict `confidence > median` visualization rule were controlled. There was one run per condition, so timings are descriptive and have no error bars.

## Overlap-Aware Frame Selection
Uniform endpoints `[0,43]` increased sequence coverage but shared little visible structure. The selected hybrid method combined camera-center distance, viewing-direction and relative-rotation angles, ORB correspondence counts [3], and fundamental-matrix RANSAC inliers [4]. It froze nested subsets satisfying `S2 ⊂ S4 ⊂ S6 ⊂ S8 ⊂ S10`. Nesting reduces frame-selection confounding because every larger condition retains all evidence from the smaller condition.

## Experimental Results
All ten configurations produced finite outputs. Allocated VRAM rose from 5.261 GiB at S2 to 6.599 GiB at S10 in both scenes; reserved memory approached 9.23 GiB. Synchronized inference remained between 0.713 and 1.142 seconds, whereas total processing reached 17.386 seconds for delivery_area and 18.645 seconds for courtyard at S10. The divergence is explained by CPU transfer, decoding, validation, visualization, and serialization scaling with dense per-view outputs.

Confidence increased overall but not monotonically. Delivery depth confidence dipped at S8 before recovering at S10; courtyard peaked at S8 and softened at S10. Courtyard confidence was higher at matched counts, but confidence is model output - not calibrated accuracy - and cannot establish that one scene was easier or reconstructed more accurately. Predicted camera-path extent also changed non-monotonically in arbitrary unaligned units.

## Qualitative Results
Shared-first-view depth retained the large door plane and ceiling structure in delivery_area and the facade, windows, foreground furniture, and ground in courtyard. Confidence concentrated around recognizable edges and structures, while point-cloud previews preserved dominant surfaces across counts. Visible differences became smaller near S8-S10. These are structured visual observations, not metric completeness or accuracy measurements.

## Strengths and Failure Modes
Observed strengths include reliable operation of all heads, coherent dominant structure, moderate allocated-memory growth, improved common visibility from overlap-aware inputs, and generally higher confidence with more views. Difficult behavior includes non-monotonic confidence, changing camera extent, diminishing high-view gains, and sensitivity of qualitative comparisons to auto-fit framing. No catastrophic failure was observed, and none is manufactured here.

## Discussion
Near-constant inference and growing total time describe different system boundaries: the model forward is efficient for this range, while artifact production grows with views. Moderate memory growth may reflect parameter reuse, chunked decoding, and allocator behavior; this is a hypothesis rather than an isolated causal measurement. Confidence changes may reflect redundant, useful, or inconsistent added evidence. The two-scene pattern suggests practical robustness, but not broad generalization.

## Limitations
The study covers two scenes and one local trajectory window per scene, with one timing run per condition. It includes no alignment, pose/depth/scan error, alternative model, order sensitivity, degradation study, robustness matrix, or fine-tuning. Confidence is not accuracy. Geometry has arbitrary scale, and independently auto-fitted previews can exaggerate apparent size differences. These boundaries prevent claims of benchmark accuracy, statistical significance, or superiority.

## Future Work
The priority is validated similarity alignment to ETH3D, followed by pose and depth or point-to-scan evaluation. Further work should add scenes, repeated timing trials, input-order tests, degradations, alternative selection strategies, and a reconstruction baseline. Fine-tuning should be considered only after aligned evaluation identifies a repeated domain-specific failure and suitable labels and resources exist.

## Conclusion
VGGT was successfully reproduced and operated reliably on consumer hardware. Overlap-aware nested sampling enabled a controlled two-scene 2/4/6/8/10-view study. Additional views generally increased confidence and visible completeness, but gains were non-monotonic and appeared to diminish at high counts. No additional inference is required for this report redesign; aligned quantitative evaluation is the appropriate next scientific phase.

## References
[1] J. Wang et al., “VGGT: Visual Geometry Grounded Transformer,” CVPR, 2025.  
[2] T. Schöps et al., “A Multi-View Stereo Benchmark with High-Resolution Images and Multi-Camera Videos,” CVPR, 2017.  
[3] E. Rublee et al., “ORB: An Efficient Alternative to SIFT or SURF,” ICCV, 2011.  
[4] M. A. Fischler and R. C. Bolles, “Random Sample Consensus,” Communications of the ACM, 1981.  
[5] M. Oquab et al., “DINOv2: Learning Robust Visual Features without Supervision,” Transactions on Machine Learning Research, 2024.
"""
    (V2/'vggt_seminar_report_v2.md').write_text(text,encoding='utf-8',newline='\n'); return text

def set_cell(cell,text,bold=False,fill=None,size=8.5):
    cell.text=''; p=cell.paragraphs[0]; r=p.add_run(str(text)); r.bold=bold; r.font.name='Arial'; r.font.size=Pt(size); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill: shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); cell._tc.get_or_add_tcPr().append(shd)
def add_table(doc,headers,rows,widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.autofit=False
    for i,h in enumerate(headers): set_cell(t.rows[0].cells[i],h,True,'DDEAF2',8)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row): set_cell(cells[i],v,False,None,7.8)
    for row in t.rows:
        for i,c in enumerate(row.cells): c.width=Inches((widths or [6.8/len(headers)]*len(headers))[i])
    return t
def page_num(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.RIGHT; run=paragraph.add_run(); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); run._r.addnext(fld)
def setup_doc():
    doc=Document(); sec=doc.sections[0]; sec.page_height=Inches(11.69); sec.page_width=Inches(8.27); sec.top_margin=sec.bottom_margin=Inches(.62); sec.left_margin=sec.right_margin=Inches(.68); sec.header_distance=sec.footer_distance=Inches(.3)
    normal=doc.styles['Normal']; normal.font.name='Arial'; normal.font.size=Pt(9.5); normal.font.color.rgb=RGBColor.from_string('1C2733'); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.15
    for name,size,color,before,after in [('Heading 1',19,'14283D',12,7),('Heading 2',13,'2B6CB0',9,5),('Heading 3',11,'20A4B8',6,3)]:
        s=doc.styles[name]; s.font.name='Arial'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    h=sec.header.paragraphs[0]; h.text='VGGT · CONTROLLED VIEW-COUNT STUDY'; h.runs[0].font.name='Arial'; h.runs[0].font.size=Pt(8); h.runs[0].font.color.rgb=RGBColor.from_string('6A7886'); page_num(sec.footer.paragraphs[0]); return doc
def pic(doc,path,width=6.9):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(4); p.add_run().add_picture(str(path),width=Inches(width))
def caption(doc,text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(7); r=p.add_run(text); r.italic=True; r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string('5D6B78')
def callout(doc,title,text,color='EEF4F8'):
    t=doc.add_table(rows=1,cols=1); set_cell(t.cell(0,0),'',False,color,9); p=t.cell(0,0).paragraphs[0]; r=p.add_run(title.upper()+'\n'); r.bold=True; r.font.color.rgb=RGBColor.from_string('2B6CB0'); r=p.add_run(text); r.font.name='Arial'; r.font.size=Pt(9.5); t.cell(0,0).width=Inches(6.8)
def newpage(doc): doc.add_page_break()
def add_text(doc,text):
    for para in text.strip().split('\n\n'):
        p=doc.add_paragraph(para); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

def main_doc(text,rows):
    doc=setup_doc()
    for _ in range(2): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('REPRODUCING VGGT'); r.font.name='Arial'; r.font.size=Pt(31); r.bold=True; r.font.color.rgb=RGBColor.from_string('14283D')
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('A Controlled Study of View-Count Scaling\nwith Overlap-Aware ETH3D Inputs'); r.font.name='Arial'; r.font.size=Pt(18); r.bold=True; r.font.color.rgb=RGBColor.from_string('2B6CB0')
    p=doc.add_paragraph('Reproduction, experimental methodology, and qualitative multi-view analysis'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    pic(doc,DIA/'hero_visual.png',6.8)
    p=doc.add_paragraph('Peleg Shpitzer   ·   Course / Instructor   ·   University of Haifa   ·   July 2026'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sections={}; current=None
    for line in text.splitlines():
        if line.startswith('## '): current=line[3:]; sections[current]=[]
        elif current and line.strip() and not line.startswith('#'): sections[current].append(line)
    newpage(doc); doc.add_heading('Abstract',1); add_text(doc,'\n'.join(sections['Abstract'])); callout(doc,'Interpretation Boundary','VGGT confidence is a model output, not calibrated reconstruction accuracy. Camera and point-cloud coordinates remain in arbitrary, unaligned prediction units.'); pic(doc,DIA/'study_at_a_glance.png',6.8); caption(doc,'Figure 1. Study at a Glance. Ten saved configurations; no new model execution.')
    newpage(doc); doc.add_heading('1. Introduction',1); add_text(doc,'\n'.join(sections['Introduction'])); callout(doc,'Research Question','How do runtime, memory, model confidence, predicted camera structure, and visible geometric completeness change as overlapping views increase from two to ten?'); pic(doc,DIA/'project_pipeline.png',6.65); caption(doc,'Figure 2. Project pipeline. Saved outputs support operational and qualitative analysis; no aligned accuracy evaluation was performed.')
    newpage(doc); doc.add_heading('2. VGGT in Brief',1); add_text(doc,'\n'.join(sections['VGGT in Brief'])); pic(doc,DIA/'vggt_architecture.png',6.75); caption(doc,'Figure 3. Simplified conceptual diagram created for this report based on VGGT [1].')
    newpage(doc); doc.add_heading('3. Reproduction Setup',1); add_text(doc,'\n'.join(sections['Reproduction Setup'])); add_table(doc,['Component','Verified setting'],[['Implementation','Official maintained VGGT; commit a288dd0f…'],['Checkpoint','VGGT-1B; SHA-256 d15bf50a…afe0'],['Compute','RTX 5080 · CUDA 13.0 · BF16 autocast'],['Runtime','Python 3.11.15 · PyTorch 2.13.0+cu130'],['Attention','Flash SDPA disabled'],['Boundary','Operational reproduction; not paper-wide metric replication']],[1.7,5.1]); caption(doc,'Table 1. Compact environment and reproduction summary. Full hashes appear in the supplement.')
    newpage(doc); doc.add_heading('4. Experimental Methodology',1); add_text(doc,'\n'.join(sections['Experimental Methodology'])); callout(doc,'Controlled Variables','Checkpoint · preprocessing · precision · hardware · order · postprocessing · confidence filter · visualization method'); pic(doc,FIG/'overlap_pair_comparison.png',6.8); caption(doc,'Figure 4. Endpoint separation increased sequence coverage but reduced common visibility. The overlap-aware pair retained substantially more shared scene content.')
    newpage(doc); doc.add_heading('5. Overlap-Aware Frame Selection',1); add_text(doc,'\n'.join(sections['Overlap-Aware Frame Selection'])); pic(doc,FIG/'nested_delivery_area.png',6.75); caption(doc,'Figure 5. Frozen nested delivery_area subsets. Orange borders identify views introduced at each step.')
    newpage(doc); pic(doc,FIG/'nested_courtyard.png',6.75); caption(doc,'Figure 6. Frozen nested courtyard subsets. Every larger set retains all earlier views.'); callout(doc,'Why Nesting Matters','The principal designed change is view count. Deterministic containment reduces frame-selection confounding and makes the comparison reproducible.')
    newpage(doc); doc.add_heading('6. Experimental Results',1); callout(doc,'Results Dashboard','5 view counts · 2 scenes · 10 configurations · all output tensors finite · 5.26 → 6.60 GiB allocated VRAM · approximately one-second inference · visible saturation around 8–10 views'); add_text(doc,'\n'.join(sections['Experimental Results'])); pic(doc,FIG/'quant_total_time.png',6.65); caption(doc,'Figure 7. Total processing time by view count. One run per condition; no error bars. Source: report/data/view_count_results.csv.')
    newpage(doc); pic(doc,FIG/'quant_memory.png',6.65); caption(doc,'Figure 8. Allocated and reserved VRAM. Scene curves overlap because tensor dimensions and code paths match.'); pic(doc,FIG/'quant_confidence.png',6.65); caption(doc,'Figure 9. Mean depth and point confidence. Model confidence - not calibrated accuracy.')
    newpage(doc); pic(doc,FIG/'quant_camera.png',6.65); caption(doc,'Figure 10. Predicted camera-path behavior in arbitrary unaligned prediction units.'); add_table(doc,['Scene','S2 total','S10 total','S2 alloc.','S10 alloc.','Confidence pattern'],[['delivery_area','3.208 s','17.386 s','5.261 GiB','6.599 GiB','overall rise; S8 dip'],['courtyard','4.510 s','18.645 s','5.261 GiB','6.599 GiB','rise to S8; slight S10 dip']],[1.15,1.05,1.05,1.05,1.05,1.5]); caption(doc,'Table 2. Main quantitative summary. Complete values are in the supplement.')
    newpage(doc); doc.add_heading('7. Qualitative Results',1); add_text(doc,'\n'.join(sections['Qualitative Results'])); pic(doc,FIG/'depth_delivery_area.png',6.8); caption(doc,'Figure 11. delivery_area shared-first-view depth across S2-S10. Colors are independently mapped prediction values, not metric depth.')
    newpage(doc); pic(doc,FIG/'depth_courtyard.png',6.8); caption(doc,'Figure 12. courtyard shared-first-view depth across S2-S10.'); pic(doc,FIG/'confidence_delivery_area.png',6.8); caption(doc,'Figure 13. delivery_area depth confidence. Model confidence is not calibrated against ETH3D error.')
    newpage(doc); pic(doc,FIG/'confidence_courtyard.png',6.8); caption(doc,'Figure 14. courtyard depth confidence.'); pic(doc,FIG/'output_diversity.png',6.8); caption(doc,'Figure 15. Article-like qualitative panel: representative inputs and diverse VGGT outputs.')
    newpage(doc); pic(doc,FIG/'delivery_area_animation_keyframes.png',5.25); caption(doc,'Figure 16. delivery_area S10 deterministic rotation keyframes from the saved PLY. Full GIF/MP4 supplied digitally.'); pic(doc,FIG/'courtyard_animation_keyframes.png',5.25); caption(doc,'Figure 17. courtyard S10 deterministic rotation keyframes. Normalization and viewing parameters are documented in the supplement.')
    newpage(doc); doc.add_heading('8. Strengths and Failure Modes',1); add_text(doc,'\n'.join(sections['Strengths and Failure Modes'])); add_table(doc,['Supported strength','Observed difficult behavior'],[['All output heads operated reliably','Confidence was non-monotonic'],['Dominant structure remained coherent','Camera extent changed with added views'],['Allocated-memory growth was moderate','High-view visible gains diminished'],['Overlap-aware selection retained shared content','Auto-fit complicates size comparison'],['Confidence generally rose with views','Confidence is not accuracy; scale is arbitrary']],[3.35,3.45]); caption(doc,'Table 3. Evidence-backed strengths and difficult behavior; no failure example was manufactured.')
    doc.add_heading('9. Discussion',1); add_text(doc,'\n'.join(sections['Discussion'])); doc.add_heading('10. Limitations',1); add_text(doc,'\n'.join(sections['Limitations'])); callout(doc,'Important Limitation','Two scenes, one local window per scene, and one timing run per condition do not support broad generalization or statistical claims.'); add_table(doc,['Limitation','Implication'],[['No alignment or metric errors','No quantitative reconstruction-accuracy claim'],['One run per condition','No variance, error bars, or significance tests'],['Arbitrary scale','No physical cross-scene camera comparison'],['Independent auto-fit','Visual size differences may be framing effects']],[2.4,4.4]); caption(doc,'Table 4. Highest-impact limitations. Full limitation inventory is in the supplement.')
    doc.add_heading('11. Future Work',1); add_text(doc,'\n'.join(sections['Future Work'])); doc.add_heading('12. Conclusion',1); add_text(doc,'\n'.join(sections['Conclusion'])); callout(doc,'Decision','No additional inference is required for the report redesign. The next scientific phase is validated similarity alignment and quantitative ETH3D evaluation.')
    doc.add_heading('References',1); add_text(doc,'\n'.join(sections['References'])); doc.add_paragraph('Detailed provenance, filenames, complete numeric tables, rendering parameters, artifact locations, and animation keyframes are provided in the supplementary PDF.'); pic(doc,DIA/'hero_visual.png',5.9); caption(doc,'Closing visual. The study connects overlapping RGB inputs to depth and saved reconstructed geometry without claiming aligned accuracy.')
    doc.core_properties.title='Reproducing VGGT: A Controlled Study of View-Count Scaling'; doc.core_properties.author='Peleg Shpitzer'; doc.save(V2/'vggt_seminar_report_v2.docx')

def supplementary(rows):
    doc=setup_doc(); doc.add_heading('VGGT Seminar Report · Supplementary Material',0); doc.add_paragraph('Peleg Shpitzer · University of Haifa · July 2026')
    doc.add_heading('S1. Reproducibility Record',1); add_table(doc,['Item','Exact value'],[['Repository commit','a288dd0f14786c93483e45524328726ab7b1b4ce'],['Checkpoint SHA-256','d15bf50a8615c8225ed48b51ea5cac673d82442ec0309036df555a053253afe0'],['Protocol','eth3d-overlap-aware-nested-v1'],['Source commits','delivery: 5ef68d3 · courtyard/cross-scene: ccb5487'],['Execution boundary','Saved outputs only; no model/checkpoint/CUDA execution']],[1.8,5.0])
    doc.add_heading('S2. Exact Frozen Frames',1); frame_rows=[]
    for r in rows: frame_rows.append([r['scene'],r['subset'],r['indices'],r['filenames'].replace('[','').replace(']','').replace('"','')])
    add_table(doc,['Scene','Set','Indices','Filenames'],frame_rows,[1.0,.55,1.5,3.75])
    newpage(doc); doc.add_heading('S3. Complete Runtime and Memory Results',1); add_table(doc,['Scene','Set','Views','Inference s','Total s','Allocated GiB','Reserved GiB'],[[r['scene'],r['subset'],r['view_count'],f"{float(r['inference_seconds']):.3f}",f"{float(r['total_seconds']):.3f}",f"{float(r['peak_allocated_gib']):.3f}",f"{float(r['peak_reserved_gib']):.3f}"] for r in rows],[1.2,.55,.55,1.0,1.0,1.2,1.2])
    doc.add_heading('S4. Complete Confidence and Geometry Results',1); add_table(doc,['Scene','Set','Depth mean','Depth median','Point mean','Point median','Retained','Path','Max sep.'],[[r['scene'],r['subset'],f"{float(r['depth_confidence_mean']):.3f}",f"{float(r['depth_confidence_median']):.3f}",f"{float(r['point_confidence_mean']):.3f}",f"{float(r['point_confidence_median']):.3f}",r['retained_points'],f"{float(r['camera_path_length']):.3f}",f"{float(r['maximum_camera_separation']):.3f}"] for r in rows],[1.0,.45,.72,.72,.72,.72,.8,.6,.6])
    newpage(doc); doc.add_heading('S5. Point-Cloud Rendering Parameters',1); doc.add_paragraph('Source: saved confidence-filtered ASCII PLY files. Deterministic sampling retained every fixed-stride vertex to approximately 45,000 points, followed by median centering and 97th-percentile radial normalization. Renders use a dark RGB(12,18,28) background, saved RGB point colors, two-pixel splats, perspective projection, a -12° elevation, and 80 evenly spaced azimuths. Animations are 960×640 at 10 fps for 8 seconds. Coordinates remain arbitrary and unaligned; normalization is visual, not geometric filtering intended to improve results.')
    pic(doc,FIG/'delivery_area_animation_keyframes.png',5.5); caption(doc,'Figure S1. delivery_area S10 rotation keyframes.'); pic(doc,FIG/'courtyard_animation_keyframes.png',5.5); caption(doc,'Figure S2. courtyard S10 rotation keyframes.')
    newpage(doc); doc.add_heading('S6. Artifact and Build Record',1); doc.add_paragraph('Canonical inputs are report/data/view_count_results.csv and the source manifests/summary files recorded in report/data/report_provenance.json. V2 is regenerated with scripts/build_v2_animations.py and scripts/build_phase10_report_v2.py, then exported through Microsoft Word to selectable-text PDFs. The build imports neither VGGT nor Torch and does not load a checkpoint.')
    doc.add_paragraph('Digital animation files: report/v2/animations/delivery_area_s10_rotation.gif/.mp4 and courtyard_s10_rotation.gif/.mp4. Main visual assets are under report/v2/figures and report/v2/diagrams.')
    doc.save(BUILD/'vggt_supplementary.docx')

def main():
    for p in (V2,FIG,DIA,TAB,SUP,BUILD,V2/'animations'): p.mkdir(parents=True,exist_ok=True)
    rows=list(csv.DictReader(DATA.open(encoding='utf-8',newline=''))); assert len(rows)==10
    hero(); glance(rows)
    flow_diagram(DIA/'vggt_architecture.png','VGGT IN BRIEF',[],True)
    flow_diagram(DIA/'project_pipeline.png','FROM CALIBRATED IMAGES TO CONTROLLED EVIDENCE',[("ETH3D SEQUENCE","calibrated images + poses"),("IMAGE + CALIBRATION INSPECTION","ordered, immutable inputs"),("OVERLAP PROXIES","pose geometry + ORB + RANSAC"),("HYBRID POSE + FEATURE SELECTION","deterministic trajectory window"),("NESTED VIEW SETS","S2 ⊂ S4 ⊂ S6 ⊂ S8 ⊂ S10"),("SAVED VGGT OUTPUTS","cameras · depth · confidence · points"),("BOUNDED ANALYSIS","runtime · memory · qualitative structure")])
    endpoint_figure(); nested('delivery_area',rows); nested('courtyard',rows)
    chart(rows,'time',FIG/'quant_total_time.png','TOTAL PROCESSING TIME','One run per condition · no error bars')
    chart(rows,'memory',FIG/'quant_memory.png','GPU MEMORY SCALING','Allocated and reserved VRAM · scene curves overlap')
    chart(rows,'confidence',FIG/'quant_confidence.png','MODEL CONFIDENCE','Model confidence — not calibrated accuracy')
    chart(rows,'camera',FIG/'quant_camera.png','PREDICTED CAMERA PATH','Arbitrary unaligned prediction units')
    gallery('delivery_area','depth', 'depth_delivery_area.png','DELIVERY AREA · SHARED-FIRST-VIEW DEPTH')
    gallery('courtyard','depth','depth_courtyard.png','COURTYARD · SHARED-FIRST-VIEW DEPTH')
    gallery('delivery_area','depth_confidence','confidence_delivery_area.png','DELIVERY AREA · DEPTH CONFIDENCE')
    gallery('courtyard','depth_confidence','confidence_courtyard.png','COURTYARD · DEPTH CONFIDENCE')
    output_diversity(); text=md_report(); main_doc(text,rows); supplementary(rows)
    shutil.copy2(ROOT/'report/tables/table01_environment.csv',TAB/'environment_full.csv'); shutil.copy2(DATA,TAB/'view_count_results_full.csv')
    print({'rows':10,'main_docx':str(V2/'vggt_seminar_report_v2.docx'),'supplement_docx':str(BUILD/'vggt_supplementary.docx'),'words':len(re.findall(r"\b[\w'-]+\b",text))})
if __name__=='__main__': main()
