"""Build Phase 9 report assets exclusively from saved canonical artifacts.

This module deliberately imports neither VGGT nor Torch and performs no inference.
"""
from __future__ import annotations

import csv
import hashlib
import json
import re
import shutil
import subprocess
from pathlib import Path
from typing import Any

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT=Path(__file__).resolve().parents[1]
REPORT=ROOT/"report"; DATA=REPORT/"data"; TABLES=REPORT/"tables"; FIGURES=REPORT/"figures"
PROTOCOL="eth3d-overlap-aware-nested-v1"
CHECKPOINT="d15bf50a8615c8225ed48b51ea5cac673d82442ec0309036df555a053253afe0"
SUBSETS=("S2","S4","S6","S8","S10")
FILENAMES={
 ("delivery_area","S2"):["DSC_0675.JPG","DSC_0681.JPG"],
 ("delivery_area","S4"):["DSC_0675.JPG","DSC_0678.JPG","DSC_0681.JPG","DSC_0684.JPG"],
 ("delivery_area","S6"):["DSC_0675.JPG","DSC_0678.JPG","DSC_0679.JPG","DSC_0680.JPG","DSC_0681.JPG","DSC_0684.JPG"],
 ("delivery_area","S8"):["DSC_0675.JPG","DSC_0676.JPG","DSC_0677.JPG","DSC_0678.JPG","DSC_0679.JPG","DSC_0680.JPG","DSC_0681.JPG","DSC_0684.JPG"],
 ("delivery_area","S10"):[f"DSC_{n:04d}.JPG" for n in range(675,685)],
 ("courtyard","S2"):["DSC_0286.JPG","DSC_0295.JPG"],
 ("courtyard","S4"):["DSC_0286.JPG","DSC_0290.JPG","DSC_0292.JPG","DSC_0295.JPG"],
 ("courtyard","S6"):["DSC_0286.JPG","DSC_0288.JPG","DSC_0289.JPG","DSC_0290.JPG","DSC_0292.JPG","DSC_0295.JPG"],
 ("courtyard","S8"):["DSC_0286.JPG","DSC_0288.JPG","DSC_0289.JPG","DSC_0290.JPG","DSC_0291.JPG","DSC_0292.JPG","DSC_0293.JPG","DSC_0295.JPG"],
 ("courtyard","S10"):[f"DSC_{n:04d}.JPG" for n in range(286,296)],
}


def sha256(path: Path) -> str:
    digest=hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda:handle.read(1024*1024),b""): digest.update(block)
    return digest.hexdigest()


def read_json(path: Path) -> Any: return json.loads(path.read_text(encoding="utf-8"))


def load_and_validate() -> tuple[list[dict[str,Any]],list[dict[str,Any]],list[Path]]:
    sources=[]; unified=[]
    for scene,phase in (("delivery_area","phase7_delivery_area_view_count"),("courtyard","phase8_courtyard_view_count")):
        base=ROOT/"outputs/experiments"/phase
        csv_path, json_path, manifest_path=base/"summary.csv",base/"summary.json",base/"manifest.json"
        sources.extend([csv_path,json_path,manifest_path])
        with csv_path.open(encoding="utf-8",newline="") as handle: csv_rows=list(csv.DictReader(handle))
        json_rows=read_json(json_path); manifest=read_json(manifest_path)
        if manifest["scene"]!=scene or manifest["protocol_version"]!=PROTOCOL or manifest["represented_configurations"]!=5: raise RuntimeError(f"Invalid manifest {scene}")
        if [r["subset"] for r in csv_rows]!=list(SUBSETS) or [r["subset"] for r in json_rows]!=list(SUBSETS): raise RuntimeError(f"Invalid subsets {scene}")
        for csv_row,json_row in zip(csv_rows,json_rows):
            for key in csv_row:
                if str(csv_row[key]).lower()!=str(json_row[key]).lower():
                    try:
                        if abs(float(csv_row[key])-float(json_row[key]))<1e-9: continue
                    except (ValueError,TypeError): pass
                    raise RuntimeError(f"CSV/JSON mismatch {scene} {csv_row['subset']} {key}")
            if csv_row["protocol_version"]!=PROTOCOL or csv_row["checkpoint_sha256"]!=CHECKPOINT: raise RuntimeError("Protocol/checkpoint mismatch")
            key=(scene,csv_row["subset"])
            unified.append({"scene":scene,"subset":csv_row["subset"],"view_count":int(csv_row["view_count"]),"protocol_version":PROTOCOL,
                "indices":json.loads(csv_row["indices"]),"filenames":FILENAMES[key],"inference_seconds":float(csv_row["inference_seconds"]),
                "total_seconds":float(csv_row["total_seconds"]),"peak_allocated_gib":float(csv_row["peak_allocated_gib"]),
                "peak_reserved_gib":float(csv_row["peak_reserved_gib"]),"depth_confidence_mean":float(csv_row["depth_confidence_mean"]),
                "depth_confidence_median":float(csv_row["depth_confidence_median"]),"point_confidence_mean":float(csv_row["point_confidence_mean"]),
                "point_confidence_median":float(csv_row["point_confidence_median"]),"retained_points":int(csv_row["retained_points"]),
                "retained_percentage":float(csv_row["retained_percentage"]),"camera_path_length":float(csv_row["camera_path_length"]),
                "maximum_camera_separation":float(csv_row["max_camera_separation"]),"checkpoint_sha256":CHECKPOINT,
                "source_commit":csv_row["git_commit"],"reused_result":csv_row["reused_existing_result"].lower()=="true"})
    cross_csv=ROOT/"outputs/experiments/phase8_courtyard_view_count/comparisons/delivery_area_vs_courtyard.csv"
    cross_json=ROOT/"outputs/experiments/phase8_courtyard_view_count/comparisons/delivery_area_vs_courtyard.json"
    sources.extend([cross_csv,cross_json])
    with cross_csv.open(encoding="utf-8",newline="") as handle: cross=list(csv.DictReader(handle))
    if [int(r["view_count"]) for r in cross] != [2,4,6,8,10] or len(read_json(cross_json))!=5: raise RuntimeError("Invalid cross-scene aggregate")
    return unified,cross,sources


def write_csv(path: Path,rows:list[dict[str,Any]]) -> None:
    with path.open("w",encoding="utf-8",newline="") as handle:
        writer=csv.DictWriter(handle,fieldnames=list(rows[0])); writer.writeheader()
        for row in rows: writer.writerow({k:json.dumps(v) if isinstance(v,(list,dict)) else v for k,v in row.items()})


def fmt(value:Any) -> str:
    if isinstance(value,bool): return "yes" if value else "no"
    if isinstance(value,float): return f"{value:.3f}"
    if isinstance(value,list): return ", ".join(map(str,value))
    return str(value)


def markdown_table(rows:list[dict[str,Any]],columns:list[tuple[str,str]]) -> str:
    lines=["| "+" | ".join(label for _,label in columns)+" |","|"+"|".join("---" for _ in columns)+"|"]
    lines += ["| "+" | ".join(fmt(row[key]).replace("|","/") for key,_ in columns)+" |" for row in rows]
    return "\n".join(lines)+"\n"


def table_assets(unified:list[dict[str,Any]],cross:list[dict[str,Any]]) -> list[tuple[str,str]]:
    environment=[{"setting":"Official repository commit","value":"a288dd0f14786c93483e45524328726ab7b1b4ce"},{"setting":"Checkpoint SHA-256","value":CHECKPOINT},
        {"setting":"Python / PyTorch / CUDA","value":"3.11.15 / 2.13.0+cu130 / 13.0"},{"setting":"GPU / precision","value":"RTX 5080 / BF16 autocast"},
        {"setting":"Flash SDPA","value":"Disabled"},{"setting":"Protocol","value":PROTOCOL},{"setting":"Inference mode","value":"Offline local checkpoint; one forward per condition"}]
    frozen=[{"scene":r["scene"],"subset":r["subset"],"indices":r["indices"],"filenames":"; ".join(r["filenames"])} for r in unified]
    runtime_cols=[("subset","Set"),("view_count","Views"),("inference_seconds","Inference s"),("total_seconds","Total s"),("peak_allocated_gib","Allocated GiB"),("peak_reserved_gib","Reserved GiB")]
    conf_cols=[("subset","Set"),("depth_confidence_mean","Depth mean"),("depth_confidence_median","Depth median"),("point_confidence_mean","Point mean"),("point_confidence_median","Point median"),("retained_points","Retained"),("camera_path_length","Camera path"),("maximum_camera_separation","Max separation")]
    cross_rows=[]
    for row in cross:
        cross_rows.append({"views":int(row["view_count"]),"delivery_total":float(row["delivery_area_total_seconds"]),"courtyard_total":float(row["courtyard_total_seconds"]),
            "delivery_depth":float(row["delivery_area_depth_confidence_mean"]),"courtyard_depth":float(row["courtyard_depth_confidence_mean"]),
            "delivery_point":float(row["delivery_area_point_confidence_mean"]),"courtyard_point":float(row["courtyard_point_confidence_mean"])})
    limitations=[{"limitation":"Two scenes / one local window each","implication":"No broad ETH3D or real-world generalization."},{"limitation":"One run per condition","implication":"No timing variance, error bars, or significance tests."},
        {"limitation":"No similarity alignment","implication":"Camera and geometry values remain arbitrary prediction units."},{"limitation":"No scan/depth/pose errors","implication":"Qualitative coherence and confidence are not reconstruction accuracy."},
        {"limitation":"Independent preview auto-fit","implication":"Point-cloud framing cannot support metric size comparison."},{"limitation":"No model/order/degradation baselines","implication":"No superiority or robustness conclusion."}]
    definitions=[("table01_environment",environment,[("setting","Setting"),("value","Value")]),
        ("table02_frozen_frames",frozen,[("scene","Scene"),("subset","Set"),("indices","Indices"),("filenames","Filenames")]),
        ("table03_delivery_runtime",[r for r in unified if r["scene"]=="delivery_area"],runtime_cols),
        ("table04_courtyard_runtime",[r for r in unified if r["scene"]=="courtyard"],runtime_cols),
        ("table05_delivery_confidence_geometry",[r for r in unified if r["scene"]=="delivery_area"],conf_cols),
        ("table06_courtyard_confidence_geometry",[r for r in unified if r["scene"]=="courtyard"],conf_cols),
        ("table07_cross_scene",cross_rows,[("views","Views"),("delivery_total","Delivery total s"),("courtyard_total","Courtyard total s"),("delivery_depth","Delivery depth conf"),("courtyard_depth","Courtyard depth conf"),("delivery_point","Delivery point conf"),("courtyard_point","Courtyard point conf")]),
        ("table08_limitations",limitations,[("limitation","Limitation"),("implication","Implication")])]
    generated=[]
    for index,(name,rows,columns) in enumerate(definitions,1):
        selected=[{key:row[key] for key,_ in columns} for row in rows]; write_csv(TABLES/f"{name}.csv",selected)
        md=f"### Table {index}. "+name.replace(f"table{index:02d}_","").replace("_"," ").title()+"\n\n"+markdown_table(rows,columns)
        (TABLES/f"{name}.md").write_text(md,encoding="utf-8"); generated.append((name,md))
    return generated


def font(size:int=18,bold:bool=False):
    for name in ("arialbd.ttf" if bold else "arial.ttf","DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf"):
        try:return ImageFont.truetype(name,size)
        except OSError:pass
    return ImageFont.load_default()


def line_plot(unified:list[dict[str,Any]],series:list[tuple[str,str,str]],title:str,ylabel:str,path:Path,note:str="") -> None:
    w,h=1200,720; m=(110,80,50,110); image=Image.new("RGB",(w,h),"white"); draw=ImageDraw.Draw(image)
    values=[float(r[key]) for r in unified for key,_,_ in series if r["scene"] in ("delivery_area","courtyard")]
    lo=0.0; hi=max(values)*1.08 if values else 1; colors={"#2463A6":"#2463A6","#D66B2C":"#D66B2C","#68A0CF":"#68A0CF","#E2A374":"#E2A374"}
    draw.line((m[0],m[1],m[0],h-m[3]),fill="black",width=2); draw.line((m[0],h-m[3],w-m[2],h-m[3]),fill="black",width=2)
    for tick in range(6):
        y=h-m[3]-tick*(h-m[1]-m[3])/5; val=lo+tick*(hi-lo)/5; draw.line((m[0]-5,y,w-m[2],y),fill="#E6E9ED",width=1); draw.text((20,y-10),f"{val:.2f}",font=font(16),fill="#333333")
    views=[2,4,6,8,10]
    for i,v in enumerate(views): x=m[0]+i*(w-m[0]-m[2])/(len(views)-1); draw.text((x-8,h-m[3]+12),str(v),font=font(17),fill="black")
    for key,label,color in series:
        scene="delivery_area" if "Delivery" in label else "courtyard"; rows=[r for r in unified if r["scene"]==scene]
        points=[]
        for i,row in enumerate(rows): x=m[0]+i*(w-m[0]-m[2])/(len(rows)-1); y=h-m[3]-(float(row[key])-lo)/(hi-lo)*(h-m[1]-m[3]); points.append((x,y))
        draw.line(points,fill=colors[color],width=4)
        for x,y in points: draw.ellipse((x-6,y-6,x+6,y+6),fill=colors[color])
    draw.text((m[0],25),title,font=font(28,True),fill="#183B5B"); draw.text((w//2-45,h-45),"Input views",font=font(18),fill="black"); draw.text((m[0],h-72),ylabel,font=font(16),fill="#333333")
    x=m[0]
    for _,label,color in series: draw.line((x,h-18,x+35,h-18),fill=colors[color],width=4); draw.text((x+42,h-29),label,font=font(15),fill="#222222"); x+=260
    if note: draw.text((w-520,48),note,font=font(13),fill="#555555")
    image.save(path)


def combine_vertical(paths:list[Path],labels:list[str],target:Path,title:str) -> None:
    items=[]
    for path in paths:
        with Image.open(path) as source: image=source.convert("RGB"); image.thumbnail((1500,680)); items.append(image.copy())
    canvas=Image.new("RGB",(1540,sum(i.height+50 for i in items)+70),"white"); draw=ImageDraw.Draw(canvas); draw.text((20,18),title,font=font(28,True),fill="#183B5B"); y=60
    for label,image in zip(labels,items): draw.text((20,y),label,font=font(18,True),fill="#333333"); y+=28; canvas.paste(image,(20,y)); y+=image.height+22
    canvas.save(target)


def figure_assets(unified:list[dict[str,Any]]) -> dict[str,str]:
    line_plot(unified,[("total_seconds","Delivery area","#2463A6"),("total_seconds","Courtyard","#D66B2C")],"Total processing time vs. view count","Seconds",FIGURES/"fig01_total_time.png","One run per condition; no error bars")
    # Two memory figures are stacked into one report figure.
    line_plot(unified,[("peak_allocated_gib","Delivery allocated","#2463A6"),("peak_allocated_gib","Courtyard allocated","#D66B2C")],"Peak allocated VRAM vs. view count","GiB",FIGURES/"_alloc.png")
    line_plot(unified,[("peak_reserved_gib","Delivery reserved","#2463A6"),("peak_reserved_gib","Courtyard reserved","#D66B2C")],"Peak reserved VRAM vs. view count","GiB",FIGURES/"_reserve.png")
    combine_vertical([FIGURES/"_alloc.png",FIGURES/"_reserve.png"],["Allocated","Reserved"],FIGURES/"fig02_vram.png","GPU memory scaling")
    (FIGURES/"_alloc.png").unlink(); (FIGURES/"_reserve.png").unlink()
    line_plot(unified,[("depth_confidence_mean","Delivery area","#2463A6"),("depth_confidence_mean","Courtyard","#D66B2C")],"Mean depth confidence vs. view count","Model confidence",FIGURES/"fig03_depth_confidence.png","Not calibrated accuracy")
    line_plot(unified,[("point_confidence_mean","Delivery area","#2463A6"),("point_confidence_mean","Courtyard","#D66B2C")],"Mean point confidence vs. view count","Model confidence",FIGURES/"fig04_point_confidence.png","Not calibrated accuracy")
    line_plot(unified,[("camera_path_length","Delivery area","#2463A6"),("camera_path_length","Courtyard","#D66B2C")],"Predicted camera path vs. view count","Arbitrary unaligned units",FIGURES/"fig05_camera_path.png","No metric cross-scene comparison")
    p7=ROOT/"outputs/experiments/phase7_delivery_area_view_count"; p8=ROOT/"outputs/experiments/phase8_courtyard_view_count"
    combine_vertical([p7/"contact_sheets/all_subsets.jpg",p8/"contact_sheets/all_subsets.jpg"],["Delivery area","Courtyard"],FIGURES/"fig06_contact_sheets.png","Frozen overlap-aware nested inputs")
    combine_vertical([p7/"comparisons/depth_view0.jpg",p8/"comparisons/depth_view0.jpg"],["Delivery area","Courtyard"],FIGURES/"fig07_depth_gallery.png","Shared first-view depth across view counts")
    combine_vertical([p7/"comparisons/depth_confidence_view0.jpg",p7/"comparisons/point_confidence_view0.jpg",p8/"comparisons/depth_confidence_view0.jpg",p8/"comparisons/point_confidence_view0.jpg"],["Delivery depth confidence","Delivery point confidence","Courtyard depth confidence","Courtyard point confidence"],FIGURES/"fig08_confidence_gallery.png","Model-confidence comparisons")
    combine_vertical([p7/"comparisons/camera_trajectories.jpg",p8/"comparisons/camera_trajectories.jpg"],["Delivery area","Courtyard"],FIGURES/"fig09_camera_gallery.png","Predicted camera centers - independently scaled")
    combine_vertical([p7/"point_cloud_previews/confidence_filtered.jpg",p8/"point_cloud_previews/confidence_filtered.jpg"],["Delivery area","Courtyard"],FIGURES/"fig10_filtered_points.png","Confidence-filtered point previews - independent auto-fit")
    return {f"fig{i:02d}":f"Figure {i}" for i in range(1,11)}


def set_run(run,size=11,bold=False,italic=False,color="000000"):
    run.font.name="Calibri"; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"),"Calibri"); run._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri")
    run.font.size=Pt(size); run.bold=bold; run.italic=italic; run.font.color.rgb=RGBColor.from_string(color)


def shade(cell,fill):
    tcpr=cell._tc.get_or_add_tcPr(); shd=tcpr.find(qn("w:shd")) or OxmlElement("w:shd"); shd.set(qn("w:fill"),fill); tcpr.append(shd) if shd.getparent() is None else None


def set_table_geometry(table,widths:list[int]):
    table.autofit=False; table.alignment=WD_TABLE_ALIGNMENT.LEFT; tblpr=table._tbl.tblPr
    tblw=tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw=OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"),str(sum(widths))); tblw.set(qn("w:type"),"dxa")
    ind=OxmlElement("w:tblInd"); ind.set(qn("w:w"),"120"); ind.set(qn("w:type"),"dxa"); tblpr.append(ind)
    grid=table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths: col=OxmlElement("w:gridCol"); col.set(qn("w:w"),str(width)); grid.append(col)
    for row in table.rows:
        for cell,width in zip(row.cells,widths):
            tcw=cell._tc.get_or_add_tcPr().get_or_add_tcW(); tcw.set(qn("w:w"),str(width)); tcw.set(qn("w:type"),"dxa"); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margin=cell._tc.get_or_add_tcPr().first_child_found_in("w:tcMar")
            if margin is None: margin=OxmlElement("w:tcMar"); cell._tc.get_or_add_tcPr().append(margin)
            for side,value in (("top",80),("bottom",80),("start",120),("end",120)):
                el=margin.find(qn(f"w:{side}")) or OxmlElement(f"w:{side}"); el.set(qn("w:w"),str(value)); el.set(qn("w:type"),"dxa"); margin.append(el) if el.getparent() is None else None


def add_docx_table(doc:Document,headers:list[str],rows:list[list[str]]):
    table=doc.add_table(rows=1,cols=len(headers)); table.style="Table Grid"; widths=[9360//len(headers)]*len(headers); widths[-1]+=9360-sum(widths)
    for index,text in enumerate(headers): table.cell(0,index).text=text; shade(table.cell(0,index),"E8EEF5")
    for row in rows:
        cells=table.add_row().cells
        for index,text in enumerate(row): cells[index].text=text
    set_table_geometry(table,widths)
    small=7.5 if len(headers)>6 else 8.5 if len(headers)>4 else 9.5
    for ridx,row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.05
                for run in p.runs:set_run(run,small,bold=ridx==0)
    return table


def add_page_number(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.RIGHT; run=paragraph.add_run("Page "); set_run(run,9,color="666666")
    fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); paragraph._p.append(fld)


def markdown_to_docx(markdown_path:Path,docx_path:Path):
    text=markdown_path.read_text(encoding="utf-8"); lines=text.splitlines(); doc=Document(); section=doc.sections[0]
    section.top_margin=section.bottom_margin=section.left_margin=section.right_margin=Inches(1); section.header_distance=section.footer_distance=Inches(.492)
    styles=doc.styles; normal=styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(8); normal.paragraph_format.line_spacing=1.333
    for name,size,color,before,after in (("Heading 1",16,"2E74B5",18,10),("Heading 2",13,"2E74B5",12,6),("Heading 3",12,"1F4D78",8,4)):
        style=styles[name]; style.font.name="Calibri"; style.font.size=Pt(size); style.font.bold=True; style.font.color.rgb=RGBColor.from_string(color); style.paragraph_format.space_before=Pt(before); style.paragraph_format.space_after=Pt(after); style.paragraph_format.keep_with_next=True
    header=section.header.paragraphs[0]; header.text="VGGT Seminar Report | Experimental Draft"; set_run(header.runs[0],9,color="6B7280")
    add_page_number(section.footer.paragraphs[0])
    # Editorial cover.
    for _ in range(5): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("SEMINAR REPORT"); set_run(r,11,True,color="7A5A00")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(lines[0][2:]); set_run(r,28,True,color="203748"); p.paragraph_format.space_after=Pt(10)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Reproduction, overlap-aware methodology, and a two-scene view-count study"); set_run(r,14,color="2B5163")
    for _ in range(4): doc.add_paragraph()
    for content in ("[Student Name]","[Course / Instructor]","July 2026"):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_run(p.add_run(content),11,color="555555")
    doc.add_page_break()
    index=next(i for i,line in enumerate(lines) if line == "## Abstract")
    def add_inline(paragraph, content):
        for part in re.split(r"(\*\*.+?\*\*|`.+?`)", content):
            if part.startswith("**") and part.endswith("**"): set_run(paragraph.add_run(part[2:-2]),11,bold=True)
            elif part.startswith("`") and part.endswith("`"): set_run(paragraph.add_run(part[1:-1]),10,color="7A3E00")
            else: set_run(paragraph.add_run(part),11)
    while index<len(lines):
        line=lines[index]
        if not line.strip(): index+=1; continue
        if line.startswith("<!--"): index+=1; continue
        if line.startswith("!["):
            match=re.match(r"!\[(.+)\]\((.+)\)",line); caption,relative=match.groups(); p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(REPORT/relative),width=Inches(6.2))
            cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_run(cp.add_run(caption),9,italic=True,color="555555"); cp.paragraph_format.space_after=Pt(10); index+=1; continue
        if line.startswith("# "): index+=1; continue
        if line.startswith("### "): doc.add_heading(line[4:],level=3); index+=1; continue
        if line.startswith("## "): doc.add_heading(line[3:],level=1); index+=1; continue
        if line.startswith("| "):
            block=[]
            while index<len(lines) and lines[index].startswith("|"): block.append(lines[index]); index+=1
            parsed=[[part.strip() for part in row.strip("|").split("|")] for row in block]
            add_docx_table(doc,parsed[0],parsed[2:]); continue
        if line.startswith("- "):
            p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.left_indent=Inches(.5); p.paragraph_format.first_line_indent=Inches(-.25); p.paragraph_format.space_after=Pt(4); add_inline(p,line[2:]); index+=1; continue
        paragraph=[line]; index+=1
        while index<len(lines) and lines[index].strip() and not lines[index].startswith(("#","|","![","- ")): paragraph.append(lines[index]); index+=1
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(p," ".join(paragraph))
    props=doc.core_properties; props.title="Reproducing VGGT and Studying View-Count Scaling"; props.subject="Deep Learning Seminar experimental report"; props.author="[Student Name]"; props.keywords="VGGT, ETH3D, view count, multi-view geometry"
    doc.save(docx_path)


def main():
    for directory in (DATA,TABLES,FIGURES): directory.mkdir(parents=True,exist_ok=True)
    unified,cross,sources=load_and_validate(); write_csv(DATA/"view_count_results.csv",unified); (DATA/"view_count_results.json").write_text(json.dumps(unified,indent=2)+"\n",encoding="utf-8")
    shutil.copy2(ROOT/"outputs/experiments/phase8_courtyard_view_count/comparisons/delivery_area_vs_courtyard.csv",DATA/"cross_scene_results.csv")
    generated=table_assets(unified,cross); figure_assets(unified)
    report_path=REPORT/"vggt_seminar_report.md"; original=report_path.read_text(encoding="utf-8"); prefix=original.split("<!-- GENERATED_TABLES -->")[0]+"<!-- GENERATED_TABLES -->\n\n"
    report_path.write_text(prefix+"\n".join(md for _,md in generated),encoding="utf-8")
    provenance={"schema_version":1,"protocol_version":PROTOCOL,"checkpoint_sha256":CHECKPOINT,"source_commits":["5ef68d3","ccb5487"],
        "sources":[{"path":str(p.relative_to(ROOT)).replace("\\","/"),"sha256":sha256(p)} for p in sources],"generated_from_saved_outputs_only":True,
        "builder":str(Path(__file__).relative_to(ROOT)).replace("\\","/"),"git_head":subprocess.run(["git","rev-parse","HEAD"],cwd=ROOT,capture_output=True,text=True,check=True).stdout.strip()}
    (DATA/"report_provenance.json").write_text(json.dumps(provenance,indent=2)+"\n",encoding="utf-8")
    markdown_to_docx(report_path,REPORT/"vggt_seminar_report.docx")
    words=len(re.findall(r"\b[\w'-]+\b",report_path.read_text(encoding="utf-8")))
    print(json.dumps({"rows":len(unified),"tables":len(generated),"figures":10,"word_count":words,"docx":str(REPORT/"vggt_seminar_report.docx")},indent=2))


if __name__=="__main__": main()
