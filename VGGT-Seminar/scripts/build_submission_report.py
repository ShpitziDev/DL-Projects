from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "report" / "submission"
STEM = "Peleg_Shpitzer_Razi_Mreeh_VGGT_Seminar_Report"
SOURCE = OUT / f"{STEM}.md"
TARGET = OUT / f"{STEM}.docx"
FIGURES = OUT / "figures"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(92, 103, 115)
TEXT = RGBColor(25, 31, 36)
GOLD = RGBColor(166, 113, 24)
LIGHT = "F2F4F7"

EQUATIONS = {
    "x_c = R x_w + t": "eq_camera_mapping.png",
    "C = −Rᵀt": "eq_camera_center.png",
    "Ĉ_i = s Q C_i + b": "eq_sim3.png",
    "α = median(d_gt) / median(d_pred), d̂ = α d_pred": "eq_depth_scale.png",
    "AbsRel = (1/N) Σ |d̂_i − d_i| / d_i": "eq_absrel.png",
    "RMSE = √[(1/N) Σ(d̂_i − d_i)²]": "eq_rmse.png",
    "ρ(confidence, absolute depth error)": "eq_spearman.png",
}


def set_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.italic = italic


def field(paragraph, code):
    run = paragraph.add_run()
    for kind, text in (("begin", None), (None, code), ("separate", None), (None, "1"), ("end", None)):
        if kind:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), kind)
        else:
            node = OxmlElement("w:instrText" if text == code else "w:t")
            node.text = text
        run._r.append(node)
    set_font(run, 8.5, color=MUTED)


def configure(doc):
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size, normal.font.color.rgb = "Calibri", Pt(11), TEXT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, NAVY),
    ):
        style = doc.styles[name]
        style.font.name, style.font.size, style.font.bold = "Calibri", Pt(size), True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name, style.font.size = "Calibri", Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
    header = section.header.paragraphs[0]
    header.text = "VGGT SEMINAR REPORT  |  CONTROLLED EVALUATION"
    for run in header.runs:
        set_font(run, 8.5, True, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    field(footer, "PAGE")


def inline(paragraph, text):
    for part in re.split(r"(\*\*.*?\*\*|\*.*?\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, italic=True)
        else:
            run = paragraph.add_run(part)
            set_font(run)


def shade(cell):
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), LIGHT)
    cell._tc.get_or_add_tcPr().append(node)


def cell_width(cell, width):
    props = cell._tc.get_or_add_tcPr()
    node = props.first_child_found_in("w:tcW")
    if node is None:
        node = OxmlElement("w:tcW")
        props.append(node)
    node.set(qn("w:w"), str(width))
    node.set(qn("w:type"), "dxa")


def add_table(doc, rows):
    values = [[x.strip() for x in row.strip().strip("|").split("|")] for row in rows]
    headers, body = values[0], values[2:]
    n = len(headers)
    if n == 2:
        widths = [2700, 6660]
    elif n == 4:
        widths = [1900, 2700, 2360, 2400]
    else:
        widths = [9360 // n] * n
        widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=1, cols=n)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
        shade(cell)
        for run in cell.paragraphs[0].runs:
            set_font(run, 8.5, True, NAVY)
    for data in body:
        cells = table.add_row().cells
        for cell, value in zip(cells, data):
            cell.text = value
            for run in cell.paragraphs[0].runs:
                set_font(run, 8.3)
    props = table._tbl.tblPr
    tbl_w = props.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    props.append(indent)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))


def add_picture(doc, name, width):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(FIGURES / name), width=Inches(width))


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    inline(p, text.strip("*"))
    for run in p.runs:
        set_font(run, 9, color=MUTED, italic=True)


def body(doc, lines):
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line == "<!-- PAGE BREAK -->":
            doc.add_page_break()
            i += 1
            continue
        image = re.fullmatch(r"!\[(.*?)\]\(figures/(.*?)\)", line)
        if image:
            name = image.group(2)
            width = 6.25
            if name in {"vggt_architecture_corrected.png", "experimental_ladder.png"}:
                width = 6.0
            elif name.startswith("eth3d_"):
                width = 6.35
            elif name == "eth3d_overlap_pair_comparison.png":
                width = 6.15
            elif name in {"p000_confidence.png", "runtime_memory.png"}:
                width = 5.15
            elif name == "eq_rotation_error.png":
                width = 6.3
            elif name == "eq_delta.png":
                width = 4.8
            else:
                width = 5.9
            add_picture(doc, name, width)
            i += 1
            if i < len(lines) and lines[i].strip().startswith("*Figure"):
                caption(doc, lines[i].strip())
                i += 1
            continue
        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i])
                i += 1
            add_table(doc, rows)
            if i < len(lines) and lines[i].strip().startswith("*Table"):
                caption(doc, lines[i].strip())
                i += 1
            continue
        heading = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading:
            p = doc.add_paragraph(style=f"Heading {min(len(heading.group(1)), 3)}")
            inline(p, heading.group(2))
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            inline(p, line[2:])
            i += 1
            continue
        stripped = line[2:-2] if line.startswith("**") and line.endswith("**") else None
        if stripped in EQUATIONS:
            add_picture(doc, EQUATIONS[stripped], 4.6)
            i += 1
            continue
        chunks = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if (not nxt or nxt.startswith("#") or nxt.startswith("|") or nxt.startswith("![")
                    or nxt.startswith("- ") or nxt == "<!-- PAGE BREAK -->"
                    or re.match(r"^\d+\.\s+", nxt)):
                break
            chunks.append(nxt)
            i += 1
        text = " ".join(chunks)
        p = doc.add_paragraph()
        if re.match(r"^\d+\.\s", text):
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
        inline(p, text)


def main():
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    configure(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(105)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("VGGT SEMINAR REPORT")
    set_font(r, 11, True, GOLD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Reproducing and Evaluating VGGT:\nView-Count Scaling, Confidence Calibration,\nand Bounded Domain Adaptation")
    set_font(r, 26, True, NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(66)
    r = p.add_run("A Controlled Study on ETH3D and TartanAir")
    set_font(r, 14, color=BLUE)
    for author in ("Peleg Shpitzer", "Razi Mreeh"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(author)
        set_font(r, 14, True, NAVY)
    body(doc, lines[lines.index("<!-- PAGE BREAK -->"):])
    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
