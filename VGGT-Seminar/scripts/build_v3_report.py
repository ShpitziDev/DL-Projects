from __future__ import annotations

import csv
import json
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "report/v3"
BASE = ROOT / "outputs/experiments/v3_tartanair_pretrained_20260729_sdpa_corrected"
ADAPTED = ROOT / "outputs/experiments/v3_tartanair_adapted_step15_20260729"
TRAIN = ROOT / "outputs/experiments/v3_tartanair_finetune_20260729"
AUDIT = ROOT / "outputs/analysis/v3_tartanair_audit_20260729"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(92, 103, 115)
LIGHT = "F2F4F7"
GOLD = RGBColor(166, 113, 24)
RED = RGBColor(155, 28, 28)


def rows(path):
    with path.open(newline="") as stream:
        return list(csv.DictReader(stream))


def set_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    tc_pr.append(element)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_font(run, 9)


def add_table(doc, headers, data, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        shade(cell, LIGHT)
        for run in cell.paragraphs[0].runs:
            set_font(run, 9, True, NAVY)
    for values in data:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = str(value)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_picture(doc, path, width=6.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, 9, color=MUTED, italic=True)


def make_figures(base, adapted):
    figures = OUT / "figures"
    figures.mkdir(parents=True, exist_ok=True)
    views = [int(r["views"]) for r in base]
    for metric, ylabel, filename in [
        ("depth_abs_rel", "Scale-aligned depth AbsRel", "depth_absrel.png"),
        ("rotation_mean_deg", "Mean camera rotation error (deg)", "camera_rotation.png"),
        ("confidence_error_spearman", "Spearman(confidence, error)", "confidence_calibration.png"),
    ]:
        plt.figure(figsize=(7.2, 3.6))
        plt.plot(views, [float(r[metric]) for r in base], "o-", lw=2.4, label="Official checkpoint")
        plt.plot(views, [float(r[metric]) for r in adapted], "s--", lw=2.4, label="Adapted heads (step 15)")
        if metric == "confidence_error_spearman":
            plt.axhline(0, color="black", lw=0.8)
        plt.xlabel("Input views")
        plt.ylabel(ylabel)
        plt.grid(alpha=0.25)
        plt.legend(frameon=False)
        plt.tight_layout()
        plt.savefig(figures / filename, dpi=180)
        plt.close()
    history = rows(TRAIN / "history.csv")
    validation = [(int(r["step"]), float(r["validation_objective"])) for r in history if r.get("validation_objective")]
    plt.figure(figsize=(7.2, 3.6))
    plt.plot([int(r["step"]) for r in history], [float(r["objective"]) for r in history],
             color="#78909C", alpha=0.7, label="Training sample objective")
    plt.plot([x for x, _ in validation], [y for _, y in validation], "o-", color="#A67118",
             lw=2.4, label="P006 validation objective")
    plt.axvline(15, color="#9B1C1C", ls="--", label="Selected step 15")
    plt.xlabel("Optimizer step")
    plt.ylabel("Adaptation objective")
    plt.grid(alpha=0.25)
    plt.legend(frameon=False)
    plt.tight_layout()
    plt.savefig(figures / "finetune_curve.png", dpi=180)
    plt.close()
    return figures


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(25, 31, 36)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, before, after, color in [
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, NAVY),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "VGGT SEMINAR STUDY  |  VERSION 3"
        for run in header.runs:
            set_font(run, 8.5, True, MUTED)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("Peleg Shpitzer & Razi Mreeh  |  University of Haifa")
        set_font(run, 8.5, color=MUTED)


def main():
    if OUT.exists():
        existing = OUT / "vggt_seminar_report_v3.docx"
        if existing.exists():
            raise FileExistsError(existing)
    OUT.mkdir(parents=True, exist_ok=True)
    base, adapted = rows(BASE / "metrics.csv"), rows(ADAPTED / "metrics.csv")
    figures = make_figures(base, adapted)
    probe = json.loads((TRAIN / "memory_probe.json").read_text())
    runtime = json.loads((TRAIN / "runtime.json").read_text())
    doc = Document()
    configure(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(96)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONTROLLED VGGT STUDY")
    set_font(r, 11, True, GOLD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("From Real-Scene Robustness to\nSynthetic Ground-Truth Adaptation")
    set_font(r, 28, True, NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Version 3 seminar report | VGGT, CVPR 2025")
    set_font(r, 14, color=BLUE)
    p.paragraph_format.space_after = Pt(64)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Peleg Shpitzer  |  Razi Mreeh")
    set_font(r, 14, True, NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("University of Haifa | July 2026")
    set_font(r, 11, color=MUTED)
    doc.add_page_break()

    add_heading(doc, "Abstract")
    doc.add_paragraph(
        "We reproduce the maintained official VGGT-1B implementation and study its behavior on "
        "two calibrated ETH3D scenes and a held-out synthetic TartanAir trajectory. The real-scene "
        "study varies overlapping input views from 2 to 10. The synthetic extension adds exact "
        "camera poses and metric depth, enabling similarity-aligned camera and scale-aligned depth "
        "evaluation. On TartanAir P000, the official checkpoint achieves 3.14% depth AbsRel at two "
        "views and 6.89% at ten views, while mean camera rotation error remains below 0.77 degrees. "
        "A leakage-safe, 30-step frozen-aggregator adaptation is operationally feasible on an RTX "
        "5080 but worsens held-out AbsRel and delta-1, demonstrating over-specialization rather "
        "than general improvement. This negative result is retained as evidence about fine-tuning risk."
    )
    add_heading(doc, "Assignment brief")
    hebrew = [
        "את התרגיל הזה יש לבסס על המאמר שהצגתם.",
        "אם אתם זוג תמצאו לזה פתרון (משהו משותף לשתי העבודות או התרכזות באחת מהם).",
        "יש לנסות את המערכת על קלטים שונים ולחפש חוזקות וחולשות.",
        "אם ניתן לעשות למערכת finetuning זה גם טוב.",
        "מה שיש להגיש זה מסמך או מצגת שמתאר מה עשיתם.",
        "בהצלחה, אילן",
    ]
    for line in hebrew:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_pr = p._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
        r = p.add_run(line)
        r.font.name = "Arial"
        r._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Arial")
        r._element.get_or_add_rPr().append(OxmlElement("w:rtl"))
        r.font.size = Pt(11)
    add_table(
        doc,
        ["Requirement", "Evidence in this submission"],
        [
            ("Based on the presented paper", "Official VGGT code/checkpoint reproduction and paper-grounded analysis"),
            ("Shared direction for a pair", "Joint focus on robustness, synthetic accuracy, and adaptation"),
            ("Different inputs", "ETH3D indoor/outdoor plus synthetic TartanAir, 2/4/6/8/10 views"),
            ("Strengths and weaknesses", "Measured geometry, confidence, memory, failure and overfitting analysis"),
            ("Fine-tuning if feasible", "Measured RTX 5080 probe, 30-step head adaptation, held-out comparison"),
            ("Document or presentation", "This illustrated report and reproducible project artifacts"),
        ],
        [2900, 6460],
    )

    add_heading(doc, "1. VGGT and the study design")
    doc.add_paragraph(
        "VGGT predicts camera parameters, depth, point maps, confidence, and tracks from one or "
        "more same-scene RGB images in a single feed-forward model. DINOv2 image tokens enter an "
        "aggregator alternating frame-local and global attention; task-specific heads decode geometry."
    )
    architecture = ROOT / "report/v2/diagrams/vggt_architecture.png"
    if architecture.exists():
        add_picture(doc, architecture, 6.1)
        add_caption(doc, "Figure 1. Project-authored summary of the VGGT architecture.")
    doc.add_paragraph(
        "The experimental ladder separates operational reproduction, controlled real-scene view-count "
        "testing, held-out synthetic scoring, and bounded adaptation. Predictions are never compared "
        "in raw scale: camera centers use Sim(3) alignment and depth uses a global median scale."
    )

    add_heading(doc, "2. Inputs and frozen protocol")
    doc.add_paragraph(
        "The existing study uses ETH3D delivery_area and courtyard with deterministic overlap-aware "
        "nested subsets. V3 adds TartanAir V2 ArchVizTinyHouseDay/easy. TartanAir is not listed in "
        "the VGGT training inventory recorded from the paper, although unreported overlap cannot be excluded."
    )
    add_picture(doc, AUDIT / "selected_contact_sheet.jpg", 6.45)
    add_caption(
        doc,
        "Figure 2. Frozen P000 frames 20-29. Adjacent pairs retain at least 136 RANSAC inliers "
        "across 1.91 m of ground-truth path motion.",
    )
    add_table(
        doc,
        ["Role", "Trajectories", "Use"],
        [
            ("Training", "P001-P005", "Only for the bounded adaptation"),
            ("Validation", "P006", "Checkpoint selection; best at step 15"),
            ("Test", "P000", "Untouched baseline/adapted comparison"),
        ],
        [1700, 2200, 5460],
    )

    add_heading(doc, "3. Pretrained synthetic baseline")
    baseline_table = []
    for row in base:
        baseline_table.append([
            f"S{row['views']}",
            f"{100*float(row['depth_abs_rel']):.2f}%",
            f"{float(row['depth_delta1']):.3f}",
            f"{float(row['rotation_mean_deg']):.3f}",
            f"{float(row['confidence_error_spearman']):+.3f}",
            f"{float(row['inference_seconds']):.3f}s",
            f"{float(row['peak_allocated_gib']):.2f}",
        ])
    add_table(
        doc,
        ["Views", "AbsRel", "delta-1", "Rot. deg", "Conf/error rho", "Forward", "GiB"],
        baseline_table,
        [850, 1200, 1100, 1100, 1700, 1300, 1110],
    )
    doc.add_paragraph(
        "The checkpoint is highly accurate for this short indoor window, but adding views does not "
        "monotonically help. Depth AbsRel more than doubles from S2 to S10, while camera rotation "
        "remains below one degree. Confidence is useful when its Spearman correlation with error is "
        "negative; at S10 the sign becomes slightly positive, exposing a calibration weakness."
    )
    add_picture(doc, figures / "depth_absrel.png", 6.25)
    add_caption(doc, "Figure 3. Held-out scale-aligned depth error before and after adaptation.")
    add_picture(doc, figures / "confidence_calibration.png", 6.25)
    add_caption(doc, "Figure 4. Confidence calibration weakens as more views are added.")

    add_heading(doc, "4. Fine-tuning feasibility and execution")
    doc.add_paragraph(
        "The maintained repository's training code is a post-publication reimplementation. Our bounded "
        "experiment freezes the 1.2B aggregator and updates only camera and depth heads using adjacent "
        "two-view samples. It uses BF16, AdamW, gradient clipping, 30 optimizer steps, and P006-only "
        "checkpoint selection. This is domain adaptation, not reproduction of the paper's 64-A100 training."
    )
    add_table(
        doc,
        ["Feasibility item", "Measured result"],
        [
            ("GPU", "NVIDIA GeForce RTX 5080, 15.92 GiB"),
            ("Forward/backward probe", f"{probe['peak_allocated_gib']:.2f} GiB allocated; {probe['peak_reserved_gib']:.2f} GiB reserved"),
            ("Optimizer steps", "30; camera and depth heads only"),
            ("Elapsed adaptation time", f"{runtime['elapsed_seconds']:.1f} seconds"),
            ("Checkpoint selection", "P006 validation minimum at step 15"),
            ("Test isolation", "P000 never used for training or selection"),
        ],
        [3000, 6360],
    )
    add_picture(doc, figures / "finetune_curve.png", 6.25)
    add_caption(doc, "Figure 5. Validation improves until step 15, then begins to regress.")

    add_heading(doc, "5. Pretrained versus adapted")
    comparison = []
    for b, a in zip(base, adapted):
        comparison.append([
            f"S{b['views']}",
            f"{100*float(b['depth_abs_rel']):.2f}%",
            f"{100*float(a['depth_abs_rel']):.2f}%",
            f"{float(b['depth_rmse']):.3f}",
            f"{float(a['depth_rmse']):.3f}",
            f"{float(b['rotation_mean_deg']):.3f}",
            f"{float(a['rotation_mean_deg']):.3f}",
        ])
    add_table(
        doc,
        ["Views", "Base AbsRel", "Adapt AbsRel", "Base RMSE", "Adapt RMSE", "Base rot.", "Adapt rot."],
        comparison,
        [800, 1400, 1450, 1400, 1450, 1430, 1430],
    )
    doc.add_paragraph(
        "Adaptation is not a held-out success. RMSE falls slightly at every view count, but AbsRel and "
        "delta-1 worsen, and camera errors rise modestly. At S10, AbsRel increases from 6.89% to "
        "12.79%. The most defensible interpretation is that a tiny two-view adaptation changes the "
        "error distribution and over-specializes to the training/validation trajectories. The official "
        "checkpoint remains the recommended model."
    )
    add_picture(doc, figures / "camera_rotation.png", 6.25)
    add_caption(doc, "Figure 6. Camera rotation error remains low but does not improve after adaptation.")

    add_heading(doc, "6. Strengths, weaknesses, and lessons")
    add_heading(doc, "Observed strengths", 2)
    for text in [
        "All official output heads operated reliably in the earlier reproduction and real-scene pilots.",
        "On synthetic P000, camera orientation stays below 0.77 degrees mean error through ten views.",
        "Two-view scale-aligned depth is strong: 3.14% AbsRel and 97.8% delta-1.",
        "Frozen-head adaptation is practical on a consumer RTX 5080 with substantial memory margin.",
    ]:
        doc.add_paragraph(text, style="List Bullet")
    add_heading(doc, "Observed weaknesses", 2)
    for text in [
        "More overlapping views do not guarantee lower depth error; S10 is worse than S2.",
        "Confidence calibration becomes unreliable at S10 even when outputs remain visually plausible.",
        "Two-point Sim(3) alignment makes S2 camera ATE degenerate; rotation and larger subsets are more informative.",
        "Short domain adaptation can improve one metric while harming more interpretable relative metrics.",
    ]:
        doc.add_paragraph(text, style="List Bullet")
    add_heading(doc, "Methodological lesson", 2)
    doc.add_paragraph(
        "A fine-tuning run should not be labeled successful because its training loss decreases. A "
        "held-out split, frozen baseline, multiple metrics, and retained negative results are essential."
    )

    add_heading(doc, "7. Limitations")
    doc.add_paragraph(
        "The quantitative synthetic study covers one environment and one ten-frame test window. "
        "TartanAir is absent from the documented VGGT training mix, but unreported overlap cannot be "
        "excluded. Depth validity is limited to 0.1-100 m to remove synthetic far-value sentinels. "
        "The fine-tune uses only 30 steps and adjacent two-view samples; it does not establish the best "
        "possible adaptation recipe. Timings are single synchronized forwards, not repeated benchmarks. "
        "The original ETH3D study remains qualitative because laser alignment was outside its frozen scope."
    )

    add_heading(doc, "8. Conclusion")
    doc.add_paragraph(
        "VGGT is operationally strong and geometrically accurate on the selected synthetic sequence, "
        "but view-count gains are not monotonic and confidence can cease to track error. Fine-tuning "
        "the camera/depth heads is feasible on an RTX 5080, yet the measured adaptation overfits and "
        "should not replace the official checkpoint. The project therefore satisfies the assignment "
        "through varied inputs, explicit strengths and weaknesses, and an honest measured fine-tuning study."
    )
    add_heading(doc, "References")
    references = [
        "Wang et al. VGGT: Visual Geometry Grounded Transformer. CVPR, 2025.",
        "Schops et al. A Multi-View Stereo Benchmark with High-Resolution Images and Multi-Camera Videos. CVPR, 2017.",
        "Wang et al. TartanAir: A Dataset to Push the Limits of Visual SLAM. IROS, 2020.",
        "Official VGGT repository, pinned project revision a288dd0f14786c93483e45524328726ab7b1b4ce.",
        "TartanAir V2 documentation and official theairlabcmu/tartanair2 distribution.",
    ]
    for text in references:
        doc.add_paragraph(text, style="List Number")
    path = OUT / "vggt_seminar_report_v3.docx"
    doc.save(path)
    (OUT / "README.md").write_text(
        "# Version 3 report\n\n"
        "Authors: Peleg Shpitzer and Razi Mreeh.\n\n"
        "Built from frozen saved results; report generation performs no inference or training.\n"
    )
    print(path)


if __name__ == "__main__":
    main()
