from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "report" / "v4"
SOURCE = OUT / "seminar_report_working.md"
TARGET = OUT / "seminar_report_working.docx"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(92, 103, 115)
TEXT = RGBColor(25, 31, 36)
GOLD = RGBColor(166, 113, 24)
LIGHT = "F2F4F7"


def font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.italic = italic


def field(paragraph, code):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    font(run, 8.5, color=MUTED)


def configure(doc):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = Inches(1)
    sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size, normal.font.color.rgb = "Calibri", Pt(11), TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    bullets = doc.styles["List Bullet"]
    bullets.font.name, bullets.font.size = "Calibri", Pt(10.5)
    bullets.paragraph_format.space_after = Pt(2)
    bullets.paragraph_format.line_spacing = 1.05
    for name, size, before, after, color in [
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, NAVY),
    ]:
        style = doc.styles[name]
        style.font.name, style.font.size, style.font.bold = "Calibri", Pt(size), True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = sec.header.paragraphs[0]
    header.text = "VGGT SEMINAR STUDY  |  CONTROLLED EVALUATION"
    for run in header.runs:
        font(run, 8.5, True, MUTED)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Peleg Shpitzer & Razi Mreeh  |  ")
    font(r, 8.5, color=MUTED)
    field(footer, "PAGE")


def inline(paragraph, text):
    # Minimal Markdown emphasis; equations and Unicode remain literal and editable.
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            font(r, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1])
            font(r, italic=True)
        else:
            r = paragraph.add_run(part)
            font(r)


def shade(cell, fill=LIGHT):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_width(cell, dxa):
    pr = cell._tc.get_or_add_tcPr()
    width = pr.first_child_found_in("w:tcW")
    if width is None:
        width = OxmlElement("w:tcW")
        pr.append(width)
    width.set(qn("w:w"), str(dxa))
    width.set(qn("w:type"), "dxa")


def add_table(doc, md_rows):
    values = [[x.strip() for x in line.strip().strip("|").split("|")] for line in md_rows]
    headers, data = values[0], values[2:]
    n = len(headers)
    total = 9360
    # Give the descriptive first column more space.
    if n <= 4:
        widths = [int(total * .34)] + [int(total * .66 / (n - 1))] * (n - 1)
    else:
        widths = [int(total / n)] * n
    widths[-1] += total - sum(widths)
    table = doc.add_table(rows=1, cols=n)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
        shade(cell)
        for r in cell.paragraphs[0].runs:
            font(r, 8.2, True, NAVY)
    for row in data:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
            for r in cell.paragraphs[0].runs:
                font(r, 8.2)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_picture(doc, rel, alt):
    path = OUT / rel
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    width = 4.9
    if path.name == "eth3d_two_scene_results.png":
        width = 6.2
    elif path.name in {"experimental_ladder.png", "vggt_architecture_corrected.png"}:
        width = 5.25
    elif path.name in {"eth3d_overlap_pair_comparison.png", "tartanair_nested_subsets.png"}:
        width = 5.75
    p.add_run().add_picture(str(path), width=Inches(width))


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    inline(p, text.strip("*"))
    for r in p.runs:
        font(r, 8.5, color=MUTED, italic=True)


def render_body(doc, lines):
    i = 0
    first_break = True
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line == "<!-- PAGE BREAK -->":
            # The first marker separates the editorial cover. Remaining markers
            # are semantic section boundaries in Markdown; Word paginates the
            # body continuously to avoid sparse spill pages.
            if first_break:
                doc.add_page_break()
                first_break = False
            i += 1
            continue
        image = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", line)
        if image:
            add_picture(doc, image.group(2), image.group(1))
            i += 1
            if i < len(lines) and lines[i].strip().startswith("*Figure"):
                add_caption(doc, lines[i].strip())
                i += 1
            continue
        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i])
                i += 1
            add_table(doc, rows)
            if i < len(lines) and lines[i].strip().startswith("*Table"):
                add_caption(doc, lines[i].strip())
                i += 1
            continue
        m = re.match(r"^(#{1,3})\s+(.*)$", line)
        if m:
            level = min(len(m.group(1)), 3)
            p = doc.add_paragraph(style=f"Heading {level}")
            inline(p, m.group(2))
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            inline(p, line[2:])
            i += 1
            continue
        if line.startswith(r"\["):
            chunks = [line]
            i += 1
            while i < len(lines):
                chunks.append(lines[i].rstrip())
                if lines[i].rstrip().endswith(r"\]"):
                    i += 1
                    break
                i += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(8)
            equation = " ".join(chunks).replace(r"\[", "").replace(r"\]", "")
            r = p.add_run(equation)
            font(r, 10, italic=True, color=NAVY)
            continue
        # Join ordinary Markdown lines into one paragraph.
        chunks = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if (not nxt or nxt.startswith("#") or nxt.startswith("|") or nxt.startswith("![")
                    or nxt.startswith("- ") or nxt == "<!-- PAGE BREAK -->" or nxt.startswith(r"\[")):
                break
            chunks.append(nxt)
            i += 1
        p = doc.add_paragraph()
        inline(p, " ".join(chunks))


def main():
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    configure(doc)

    # Editorial cover: consume source title block, then build it with deliberate hierarchy.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("VGGT SEMINAR STUDY")
    font(r, 11, True, GOLD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Reproducing and Evaluating VGGT:\nView-Count Scaling, Confidence Calibration,\nand Bounded Domain Adaptation")
    font(r, 25, True, NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(58)
    r = p.add_run("A controlled study on ETH3D and TartanAir")
    font(r, 14, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Peleg Shpitzer  |  Razi Mreeh")
    font(r, 14, True, NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("University of Haifa")
    font(r, 11, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Course: [course name]  |  Instructor: [instructor name]")
    font(r, 10, color=MUTED, italic=True)

    start = lines.index("<!-- PAGE BREAK -->")
    render_body(doc, lines[start:])
    TARGET.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
